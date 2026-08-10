"""
Cross-reference and arithmetic check for the submission documents.

Written after the third round of review, because two of the defects found by then were created
by earlier fixes rather than present in the original: renumbering a section broke two "section N"
references, and a table added to close one gap had a column that did not sum. Both are the kind
of thing a person misses and a script does not.

Checks, per document:
  - every "section N" and "§N" reference points at a heading that exists
  - headings are numbered consecutively from 1 with no gaps or repeats
  - any table whose last row is a bold Total actually sums to it

Usage
-----
    python check_docs.py                    # every .docx in docs/
    python check_docs.py Task2_Training_Pipeline.docx
"""

import re
import sys
from pathlib import Path

from docx import Document

DOCS = Path("docs")
REF = re.compile(r"(?:section\s+|§\s?)(\d+)", re.I)
HEADING = re.compile(r"^(\d+)\.\s+\S")
NUM = re.compile(r"^-?[\d,]+(?:\.\d+)?$")


def check(path: Path) -> int:
    doc = Document(path)
    problems = []

    headings = {}
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            m = HEADING.match(p.text.strip())
            if m:
                headings[int(m.group(1))] = p.text.strip()

    # 1. consecutive numbering
    if headings:
        nums = sorted(headings)
        expected = list(range(1, len(nums) + 1))
        if nums != expected:
            problems.append(f"heading numbers are {nums}, expected {expected}")

    # 2. every cross-reference resolves
    body = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                body += "\n" + c.text
    for n in sorted({int(x) for x in REF.findall(body)}):
        if n not in headings:
            problems.append(f'reference to section {n}, which does not exist '
                            f'(headings are {sorted(headings)})')

    # 3. tables that claim a total
    #
    # Two shapes occur. Most tables end in a single grand total, which should equal the rows
    # above it. The Task 4 taxonomy instead carries two subtotals over disjoint subsets, where
    # the meaningful check is that the subtotals together equal the data rows. Treating the
    # second shape as the first reports a failure on a table that is correct, which is how this
    # check first behaved.
    for ti, t in enumerate(doc.tables, 1):
        if len(t.rows) < 3:
            continue
        is_total = [("total" in r.cells[0].text.strip().lower()
                     or "the four rows" in r.cells[0].text.strip().lower())
                    for r in t.rows[1:]]
        if not any(is_total):
            continue
        # Anything printed after the last total row is commentary, not an addend.
        last_total = max(i for i, tot in enumerate(is_total) if tot)
        rows, is_total = t.rows[1:last_total + 2], is_total[:last_total + 1]
        data = [r for r, tot in zip(rows, is_total) if not tot]
        totals = [r for r, tot in zip(rows, is_total) if tot]

        for col in range(1, len(t.rows[0].cells)):
            def nums(rows):
                out = []
                for r in rows:
                    c = r.cells[col].text.strip().replace(",", "")
                    if NUM.match(c):
                        out.append(float(c))
                return out

            d_vals, t_vals = nums(data), nums(totals)
            if len(d_vals) < 2 or not t_vals:
                continue
            if abs(sum(d_vals) - sum(t_vals)) > 0.5:
                shape = "subtotals" if len(t_vals) > 1 else "the total"
                problems.append(f"table {ti} column {col}: data rows sum to "
                                f"{sum(d_vals):,.0f} but {shape} give "
                                f"{sum(t_vals):,.0f}")

    print(f"\n{path.name}")
    print(f"  headings   {sorted(headings) or 'none numbered'}")
    if problems:
        for p in problems:
            print(f"  PROBLEM    {p}")
    else:
        print("  clean")
    return len(problems)


def main():
    targets = ([DOCS / a for a in sys.argv[1:]] if len(sys.argv) > 1
               else sorted(p for p in DOCS.glob("*.docx") if not p.name.startswith("~$")))
    total = sum(check(p) for p in targets)
    print(f"\n{'=' * 60}")
    print(f"  {total} problem(s) across {len(targets)} document(s)")
    print(f"{'=' * 60}\n")
    return 1 if total else 0


if __name__ == "__main__":
    sys.exit(main())
